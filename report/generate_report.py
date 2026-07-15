#!/usr/bin/env python3
from __future__ import annotations

import argparse
import os
import re
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
DOCX_PATH = DIST / "final-project-report.docx"

FONT_BODY = "微软雅黑"
FONT_CODE = "Consolas"
ACCENT = "2563EB"
ACCENT_DARK = "1E3A8A"
TEXT = "111827"
MUTED = "64748B"
BORDER = "CBD5E1"
ACCENT_LIGHT = "EFF6FF"

COURSE_NAME = "Python 程序设计"
ASSIGNMENT_NAME = "期末大作业"
DOCUMENT_TITLE = "Python 程序设计期末大作业报告"
SUBTITLE_DEFAULT = "Python 期末大作业"
INFO_FIELDS = [('学生姓名', 'student.name'), ('学号', 'student.id'), ('班级', 'student.class_name'), ('项目名称', 'project.name'), ('所选题目', 'project.topic'), ('仓库地址', 'project.repository_url')]
GRADING_ROWS = [('功能完整性', 25), ('CLI 入口', 12), ('Streamlit 页面', 15), ('数据建模与 JSON 持久化', 15), ('模块化结构与代码质量', 10), ('异常处理与健壮性', 8), ('pytest 测试', 10), ('文档、截图与 AI 反思', 5)]
CONTENT_SECTIONS = [('项目设计说明', 'docs/DESIGN.md'), ('使用说明', 'docs/USAGE_EXAMPLES.md'), ('AI 使用反思', 'docs/AI_REFLECTION.md'), ('运行截图', 'docs/SCREENSHOTS.md')]
AUTO_SCREENSHOTS = True


def read_text(path: str | Path, default: str = "") -> str:
    p = Path(path)
    if not p.is_absolute():
        p = ROOT / p
    return p.read_text(encoding="utf-8") if p.exists() else default


def parse_simple_yaml(text: str) -> dict:
    data: dict[str, object] = {}
    stack: list[tuple[int, dict]] = [(-1, data)]
    for raw in text.splitlines():
        if not raw.strip() or raw.lstrip().startswith("#"):
            continue
        indent = len(raw) - len(raw.lstrip(" "))
        line = raw.strip()
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        while stack and indent <= stack[-1][0]:
            stack.pop()
        parent = stack[-1][1]
        if value == "":
            child: dict[str, object] = {}
            parent[key] = child
            stack.append((indent, child))
        else:
            parent[key] = value
    return data


def get_nested(data: dict, path: str, default: str = "") -> str:
    current: object = data
    for part in path.split("."):
        if not isinstance(current, dict) or part not in current:
            return default
        current = current[part]
    return str(current)


def current_commit() -> str:
    env_sha = os.environ.get("GITHUB_SHA") or os.environ.get("GITEA_SHA") or os.environ.get("CI_COMMIT_SHA")
    if env_sha:
        return env_sha[:12]
    try:
        return subprocess.check_output(
            ["git", "rev-parse", "--short=12", "HEAD"],
            cwd=ROOT,
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except Exception:
        return ""


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, *, font_name: str = FONT_BODY, size: int | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    set_style_font(styles["Normal"], size=10, color=TEXT)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 22, TEXT),
        ("Heading 1", 16, TEXT),
        ("Heading 2", 13, ACCENT_DARK),
        ("Heading 3", 11, TEXT),
    ]:
        set_style_font(styles[style_name], size=size, color=color, bold=True)
        styles[style_name].paragraph_format.space_before = Pt(10)
        styles[style_name].paragraph_format.space_after = Pt(6)

    code_style = styles["Code Block"] if "Code Block" in styles else styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code_style, font_name=FONT_CODE, size=9, color="334155")
    code_style.paragraph_format.left_indent = Cm(0.35)
    code_style.paragraph_format.right_indent = Cm(0.2)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = TEXT, size: float = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text or ""))
    set_east_asia_font(run, FONT_BODY)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, *, header_fill: str = ACCENT_DARK) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)


def add_run(paragraph, text: str, *, bold: bool = False, code: bool = False) -> None:
    run = paragraph.add_run(text)
    set_east_asia_font(run, FONT_CODE if code else FONT_BODY)
    run.bold = bold
    run.font.size = Pt(9 if code else 10)
    run.font.color.rgb = RGBColor.from_string("334155" if code else TEXT)


def add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            add_run(paragraph, text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            add_run(paragraph, token[2:-2], bold=True)
        else:
            add_run(paragraph, token[1:-1], code=True)
        pos = match.end()
    if pos < len(text):
        add_run(paragraph, text[pos:])


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    add_inline_runs(p, text)
    return p


def add_cover(doc: Document, info: dict) -> None:
    course = get_nested(info, "course", COURSE_NAME) or COURSE_NAME
    assignment = get_nested(info, "assignment", ASSIGNMENT_NAME) or ASSIGNMENT_NAME
    project_title = get_nested(info, "project_name", "") or get_nested(info, "project.name", "") or SUBTITLE_DEFAULT

    p = doc.add_paragraph(course, style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph(DOCUMENT_TITLE, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(project_title, style="Heading 1")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph(f"{assignment} · 自动生成提交文档")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        set_east_asia_font(run, FONT_BODY)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph("")
    rows: list[tuple[str, str]] = []
    for label, path in INFO_FIELDS:
        rows.append((label, get_nested(info, path, "")))
    rows.append(("最终提交", current_commit()))
    rows.append(("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")))

    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.4)
    table.columns[1].width = Cm(12.8)
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        set_cell_text(cells[0], label, bold=True, color=ACCENT_DARK)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], ACCENT_LIGHT)
        set_cell_border(cells[0])
        set_cell_border(cells[1])

    doc.add_paragraph("")
    tip = doc.add_paragraph("说明：请下载 Word 后检查内容、截图与分页。若需要，可做少量排版微调后打印提交。")
    tip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tip.runs:
        set_east_asia_font(run, FONT_BODY)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_grading_page(doc: Document) -> None:
    doc.add_paragraph("教师评分页", style="Heading 1")
    add_paragraph(doc, "本页用于纸质批阅或归档。评分细则以仓库中的公开评分标准为准。")

    headers = ["评分项", "分值", "得分", "备注"]
    table = doc.add_table(rows=1 + len(GRADING_ROWS) + 1, cols=4)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, (item, points) in enumerate(GRADING_ROWS, start=1):
        cells = table.rows[r_idx].cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], str(points), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], "")

    total_row = table.rows[len(GRADING_ROWS) + 1].cells
    for idx, value in enumerate(["总分", "100", "", ""]):
        set_cell_text(total_row[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER if idx in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(total_row[idx], "FEFCE8")

    style_table(table)

    doc.add_paragraph("教师评语", style="Heading 2")
    comment = doc.add_table(rows=1, cols=1)
    comment.rows[0].height = Cm(4.2)
    set_cell_text(comment.cell(0, 0), "")
    set_cell_border(comment.cell(0, 0))

    p = doc.add_paragraph()
    add_run(p, "教师签名：________________        日期：________________")
    doc.add_page_break()


def resolve_doc_path(md_path: str, src: str) -> Path:
    src = src.strip().strip('"').strip("'")
    p = Path(src)
    if p.is_absolute():
        return p
    return (ROOT / md_path).parent / p


def add_image(doc: Document, image_path: Path, alt: str = "") -> None:
    try:
        display = image_path.relative_to(ROOT)
    except ValueError:
        display = image_path
    if not image_path.exists():
        add_paragraph(doc, f"[截图缺失：{display}]")
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.7))
        caption = alt or image_path.name
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            set_east_asia_font(run, FONT_BODY)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(MUTED)
    except Exception as exc:
        add_paragraph(doc, f"[截图无法读取：{display}，{exc}]")


def markdown_to_docx(doc: Document, markdown: str, md_path: str) -> None:
    paragraph_buf: list[str] = []
    code_buf: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buf
        if paragraph_buf:
            add_paragraph(doc, " ".join(paragraph_buf))
            paragraph_buf = []

    def flush_code() -> None:
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph(style="Code Block")
            p.add_run("\n".join(code_buf))
            for run in p.runs:
                set_east_asia_font(run, FONT_CODE)
                run.font.size = Pt(9)
            code_buf = []

    for raw in markdown.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not stripped:
            flush_paragraph()
            continue
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            add_image(doc, resolve_doc_path(md_path, image_match.group(2)), image_match.group(1))
            continue
        if stripped.startswith("#"):
            flush_paragraph()
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            text = stripped[level:].strip()
            doc.add_paragraph(text, style=f"Heading {level}")
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            continue
        paragraph_buf.append(stripped)
    flush_paragraph()
    flush_code()


def drop_initial_h1(markdown: str) -> str:
    lines = markdown.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and lines[0].lstrip().startswith("# "):
        lines = lines[1:]
        if lines and not lines[0].strip():
            lines = lines[1:]
    return "\n".join(lines)


def add_content_sections(doc: Document) -> None:
    for index, (title, path) in enumerate(CONTENT_SECTIONS):
        if index > 0:
            doc.add_page_break()
        doc.add_paragraph(title, style="Heading 1")
        markdown = drop_initial_h1(read_text(path, f"# {title}\n\n待填写。"))
        markdown_to_docx(doc, markdown, path)

    if AUTO_SCREENSHOTS:
        images_dir = ROOT / "docs" / "screenshots"
        images = sorted(p for p in images_dir.glob("*") if p.suffix.lower() in {".png", ".jpg", ".jpeg"}) if images_dir.exists() else []
        if images:
            doc.add_page_break()
            doc.add_paragraph("运行截图", style="Heading 1")
            for image in images:
                add_image(doc, image, image.name)


def build_docx() -> None:
    DIST.mkdir(exist_ok=True)
    info = parse_simple_yaml(read_text("docs/STUDENT_INFO.yml"))
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = DOCUMENT_TITLE
    doc.core_properties.subject = get_nested(info, "project_name", "") or get_nested(info, "project.name", "")
    doc.core_properties.author = get_nested(info, "student.name", "")

    add_cover(doc, info)
    add_grading_page(doc)
    add_content_sections(doc)

    doc.save(DOCX_PATH)
    print(f"generated {DOCX_PATH}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", action="store_true", help="generate editable Word report")
    parser.add_argument("--pdf", action="store_true", help="deprecated; generate DOCX instead")
    parser.parse_args()
    build_docx()


if __name__ == "__main__":
    main()
